"""Unstructured file reader.

A parser for unstructured text files using Unstructured.io.
Supports .txt, .docx, .pptx, .ppt, .jpg, .png, .eml, .html, and .pdf documents.

To use .doc and .xls parser, install

sudo apt-get install -y libmagic-dev poppler-utils libreoffice
pip install xlrd

"""
import mimetypes
from pathlib import Path
from typing import Any, Dict, List, Optional
import re
import zipfile

from llama_index.core.readers.base import BaseReader

from kotaemon.base import Document


class UnstructuredReader(BaseReader):
    """General unstructured text reader for a variety of files."""

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        """Init params."""
        super().__init__(*args)  # not passing kwargs to parent bc it cannot accept it

        self.api = False  # we default to local
        if "url" in kwargs:
            self.server_url = str(kwargs["url"])
            self.api = True  # is url was set, switch to api
        else:
            self.server_url = "http://localhost:8000"

        if "api" in kwargs:
            self.api = kwargs["api"]

        self.api_key = ""
        if "api_key" in kwargs:
            self.api_key = kwargs["api_key"]

    """ Loads data using Unstructured.io

        Depending on the construction if url is set or api = True
        it'll parse file using API call, else parse it locally
        additional_metadata is extended by the returned metadata if
        split_documents is True

        Returns list of documents
    """

    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict] = None,
        split_documents: Optional[bool] = False,
        **kwargs,
    ) -> List[Document]:
        """If api is set, parse through api"""
        file_path_str = str(file)
        if self.api:
            from unstructured.partition.api import partition_via_api

            elements = partition_via_api(
                filename=file_path_str,
                api_key=self.api_key,
                api_url=self.server_url + "/general/v0/general",
            )
        else:
            """Parse file locally"""
            from unstructured.partition.auto import partition

            partition_kwargs = dict(kwargs)
            content_type = partition_kwargs.get("content_type") or self._infer_content_type(
                file_path_str, extra_info
            )
            if content_type:
                partition_kwargs["content_type"] = content_type

            metadata_filename = (
                partition_kwargs.get("metadata_filename")
                or partition_kwargs.get("file_filename")
                or self._infer_file_name(
                file_path_str, extra_info, content_type
                )
            )
            partition_kwargs.pop("file_filename", None)
            if metadata_filename:
                partition_kwargs["metadata_filename"] = metadata_filename

            try:
                elements = partition(filename=file_path_str, **partition_kwargs)
            except Exception:
                return self._fallback_load_data(
                    file_path=file,
                    extra_info=extra_info,
                    split_documents=split_documents,
                    content_type=content_type,
                )

        """ Process elements """
        docs = []
        file_name = Path(file).name
        file_path = str(Path(file).resolve())
        if split_documents:
            for node in elements:
                metadata = {"file_name": file_name, "file_path": file_path}
                if hasattr(node, "metadata"):
                    """Load metadata fields"""
                    for field, val in vars(node.metadata).items():
                        if field == "_known_field_names":
                            continue
                        # removing coordinates because it does not serialize
                        # and dont want to bother with it
                        if field == "coordinates":
                            continue
                        # removing bc it might cause interference
                        if field == "parent_id":
                            continue
                        metadata[field] = val

                if extra_info is not None:
                    metadata.update(extra_info)

                metadata["file_name"] = file_name
                docs.append(Document(text=node.text, metadata=metadata))

        else:
            text_chunks = [" ".join(str(el).split()) for el in elements]
            metadata = {"file_name": file_name, "file_path": file_path}

            if extra_info is not None:
                metadata.update(extra_info)

            # Create a single document by joining all the texts
            docs.append(Document(text="\n\n".join(text_chunks), metadata=metadata))

        return docs

    def _fallback_load_data(
        self,
        file_path: Path,
        extra_info: Optional[Dict],
        split_documents: bool,
        content_type: Optional[str],
    ) -> List[Document]:
        file_name = Path(file_path).name
        file_path_abs = str(Path(file_path).resolve())
        metadata = {"file_name": file_name, "file_path": file_path_abs}
        if extra_info is not None:
            metadata.update(extra_info)

        text = self._extract_text_without_unstructured(str(file_path), content_type)
        if not text:
            return [Document(text="", metadata=metadata)]

        if split_documents:
            chunks = [chunk.strip() for chunk in text.split("\n\n") if chunk.strip()]
            if not chunks:
                chunks = [text.strip()]
            return [Document(text=chunk, metadata=metadata) for chunk in chunks]

        return [Document(text=text, metadata=metadata)]

    @staticmethod
    def _extract_text_without_unstructured(file_path: str, content_type: Optional[str]) -> str:
        content_type = (content_type or "").lower()

        if content_type.endswith("wordprocessingml.document"):
            try:
                from docx import Document as DocxDocument

                doc = DocxDocument(file_path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
                return "\n".join(paragraphs)
            except Exception:
                pass

        try:
            if zipfile.is_zipfile(file_path):
                chunks: List[str] = []
                with zipfile.ZipFile(file_path) as zf:
                    xml_names = [
                        name
                        for name in zf.namelist()
                        if name.endswith(".xml")
                        and (
                            name.startswith("word/")
                            or name.startswith("ppt/")
                            or name.startswith("xl/")
                        )
                    ]
                    for name in xml_names:
                        raw = zf.read(name).decode("utf-8", errors="ignore")
                        plain = re.sub(r"<[^>]+>", " ", raw)
                        plain = re.sub(r"\s+", " ", plain).strip()
                        if plain:
                            chunks.append(plain)
                if chunks:
                    return "\n".join(chunks)
        except Exception:
            pass

        try:
            return Path(file_path).read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    @staticmethod
    def _infer_file_name(
        file_path: str, extra_info: Optional[Dict], content_type: Optional[str]
    ) -> str:
        candidate = ""
        if extra_info:
            for key in ("original_file_name", "source_file_name", "file_name", "filename"):
                value = extra_info.get(key)
                if isinstance(value, str) and value.strip():
                    candidate = value.strip()
                    break
        if candidate:
            return candidate

        path_name = Path(file_path).name
        if Path(path_name).suffix:
            return path_name

        ext_map = {
            "application/pdf": ".pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
            "application/vnd.ms-powerpoint": ".ppt",
        }
        inferred_ext = ext_map.get(content_type or "", "")
        return f"{path_name}{inferred_ext}" if inferred_ext else path_name

    @staticmethod
    def _infer_content_type(file_path: str, extra_info: Optional[Dict]) -> Optional[str]:
        guessed, _ = mimetypes.guess_type(file_path)
        if guessed:
            return guessed

        if extra_info:
            for key in ("original_file_name", "source_file_name", "file_name", "filename"):
                value = extra_info.get(key)
                if isinstance(value, str) and value.strip():
                    guessed, _ = mimetypes.guess_type(value.strip())
                    if guessed:
                        return guessed

        try:
            with open(file_path, "rb") as f:
                header = f.read(8)
            if header.startswith(b"%PDF"):
                return "application/pdf"
        except Exception:
            pass

        try:
            if zipfile.is_zipfile(file_path):
                with zipfile.ZipFile(file_path) as zf:
                    names = set(zf.namelist())
                if "word/document.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if "ppt/presentation.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.presentationml.presentation"
                if "xl/workbook.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except Exception:
            pass

        return None
